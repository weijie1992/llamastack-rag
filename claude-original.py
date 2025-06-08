#!/usr/bin/env python3
"""
Document Processing Pipeline for Service Request RAG Application
This script processes legacy Service Request documents and ingests them into Llama Stack
"""

import os
import asyncio
from pathlib import Path
from typing import List, Dict, Any
from datetime import datetime

# Llama Stack SDK
from llama_stack_client import LlamaStackClient

# Document processing libraries
import PyPDF2
from docx import Document as DocxDocument
import mammoth  # For better DOCX to text conversion


class ServiceRequestProcessor:
    def __init__(self, llama_stack_port: str = "8321"):
        self.llama_stack_port = llama_stack_port
        self.client = LlamaStackClient(base_url=f"http://localhost:{llama_stack_port}")
        self.chunk_size = 1000  # Characters per chunk
        self.chunk_overlap = 200  # Overlap between chunks

    def extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF files"""
        try:
            with open(file_path, "rb") as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                return text
        except Exception as e:
            print(f"Error processing PDF {file_path}: {e}")
            return ""

    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX files (modern Word format)"""
        try:
            # Using mammoth for better formatting preservation
            with open(file_path, "rb") as docx_file:
                result = mammoth.extract_raw_text(docx_file)
                return result.value
        except Exception as e:
            print(f"Error processing DOCX {file_path}: {e}")
            # Fallback to python-docx
            try:
                doc = DocxDocument(file_path)
                text = ""
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"
                return text
            except:
                return ""

    def extract_text_from_doc(self, file_path: str) -> str:
        """Extract text from DOC files (legacy Word format)"""
        try:
            # Try using python-docx (sometimes works with .doc)
            doc = DocxDocument(file_path)
            print(f"==>> doc: {doc}")
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            print(f"Error processing DOC with python-docx {file_path}: {e}")
            # Alternative: suggest using LibreOffice or external conversion
            print(
                f"⚠️  Consider converting {file_path} to .docx format for better extraction"
            )
            return ""

    def extract_text_from_txt(self, file_path: str) -> str:
        """Extract text from TXT/MD files"""
        try:
            with open(file_path, "r", encoding="utf-8") as file:
                return file.read()
        except Exception as e:
            print(f"Error processing TXT {file_path}: {e}")
            return ""

    def extract_document_text(self, file_path: str) -> str:
        """Extract text based on file extension"""
        file_ext = Path(file_path).suffix.lower()

        if file_ext == ".pdf":
            return self.extract_text_from_pdf(file_path)
        elif file_ext == ".docx":
            return self.extract_text_from_docx(file_path)
        elif file_ext == ".doc":
            return self.extract_text_from_doc(file_path)
        elif file_ext in [".txt", ".md"]:
            return self.extract_text_from_txt(file_path)
        else:
            print(f"Unsupported file type: {file_ext}")
            return ""

    def chunk_text(self, text: str, metadata: Dict[str, Any]) -> List[Dict[str, Any]]:
        """Split text into chunks for better retrieval"""
        chunks = []
        words = text.split()

        for i in range(0, len(words), self.chunk_size):
            chunk_words = words[i : i + self.chunk_size + self.chunk_overlap]
            chunk_text = " ".join(chunk_words)

            chunk_metadata = metadata.copy()
            chunk_metadata.update(
                {
                    "chunk_index": len(chunks),
                    "chunk_start": i,
                    "chunk_end": min(i + self.chunk_size, len(words)),
                }
            )

            chunks.append({"text": chunk_text, "metadata": chunk_metadata})

        return chunks

    def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract metadata from file path and content"""
        path = Path(file_path)

        # Basic metadata
        metadata = {
            "filename": path.name,
            "file_type": path.suffix.lower(),
            "file_path": str(path),
            "category": path.parent.name,  # specs, requirements, apis, etc.
            "processed_date": datetime.now().isoformat(),
            "file_size": path.stat().st_size if path.exists() else 0,
        }

        # Try to extract additional metadata from filename
        filename_lower = path.stem.lower()

        # Identify service type from filename
        if "user" in filename_lower or "auth" in filename_lower:
            metadata["service_type"] = "authentication"
        elif "payment" in filename_lower or "billing" in filename_lower:
            metadata["service_type"] = "payment"
        elif "report" in filename_lower or "analytics" in filename_lower:
            metadata["service_type"] = "reporting"
        elif "api" in filename_lower:
            metadata["service_type"] = "api"
        else:
            metadata["service_type"] = "general"

        # Identify document type
        if "spec" in filename_lower or "specification" in filename_lower:
            metadata["document_type"] = "specification"
        elif "requirement" in filename_lower:
            metadata["document_type"] = "requirements"
        elif "api" in filename_lower:
            metadata["document_type"] = "api_documentation"
        else:
            metadata["document_type"] = "general"

        return metadata

    def register_vector_store(self, vector_db_id: str = "service_requests_db"):
        """Register a vector database using the SDK"""
        try:
            response = self.client.vector_dbs.register(
                vector_db_id=vector_db_id,
                embedding_model="all-MiniLM-L6-v2",
                embedding_dimension=384,
                provider_id="faiss",
            )
            print(f"✅ Registered vector database: {response}")
            print(f"✅ Registered vector database: {vector_db_id}")
            return True
        except Exception as e:
            # Check if it already exists
            if "already exists" in str(e).lower():
                print(f"✅ Vector database '{vector_db_id}' already exists")
                return True
            else:
                print(f"❌ Failed to register vector database: {e}")
                return False

    def list_vector_stores(self):
        """List all available vector databases"""
        try:
            response = self.client.vector_dbs.list()
            print("📋 Available vector databases:")
            for db in response:
                print(f"  - {db.vector_db_id} (model: {db.embedding_model})")
            return response
        except Exception as e:
            print(f"❌ Failed to list vector databases: {e}")
            return []

    def ingest_document_chunks(
        self, chunks: List[Dict[str, Any]], vector_db_id: str = "service_requests_db"
    ):
        """Ingest document chunks into the vector database using SDK"""

        # Prepare documents for insertion
        documents = []
        for chunk in chunks:
            documents.append(
                {
                    "document_id": f"{chunk['metadata']['filename']}_{chunk['metadata']['chunk_index']}",
                    "content": chunk["text"],
                    "metadata": chunk["metadata"],
                }
            )

        try:
            response = self.client.vector_dbs.insert(
                vector_db_id=vector_db_id, documents=documents
            )
            print(f"✅ Ingested {len(documents)} chunks")
            return True
        except Exception as e:
            print(f"❌ Failed to ingest chunks: {e}")
            return False

    def process_directory(
        self, directory_path: str, vector_db_id: str = "service_requests_db"
    ):
        """Process all documents in a directory"""
        directory = Path(directory_path)

        if not directory.exists():
            print(f"❌ Directory does not exist: {directory_path}")
            return

        # Register vector database
        if not self.register_vector_store(vector_db_id):
            print(
                f"❌ Failed to register vector database. Checking existing databases..."
            )
            self.list_vector_stores()
            return

        # Supported file extensions
        supported_extensions = {".pdf", ".docx", ".doc", ".txt", ".md"}

        # Process all files
        processed_count = 0
        total_chunks = 0

        for file_path in directory.rglob("*"):
            if file_path.is_file() and file_path.suffix.lower() in supported_extensions:
                print(f"📄 Processing: {file_path.name}")

                # Extract text
                text = self.extract_document_text(str(file_path))
                if not text.strip():
                    print(f"⚠️  No text extracted from {file_path.name}")
                    continue

                # Extract metadata
                metadata = self.extract_metadata(str(file_path))

                # Create chunks
                chunks = self.chunk_text(text, metadata)

                # Ingest chunks
                success = self.ingest_document_chunks(chunks, vector_db_id)
                if success:
                    processed_count += 1
                    total_chunks += len(chunks)
                    print(f"✅ Processed {file_path.name} - {len(chunks)} chunks")
                else:
                    print(f"❌ Failed to process {file_path.name}")

        print(f"\n🎉 Processing complete!")
        print(f"📊 Files processed: {processed_count}")
        print(f"📊 Total chunks created: {total_chunks}")
        print(f"🗃️  Vector database: {vector_db_id}")


def main():
    # Initialize processor
    processor = ServiceRequestProcessor()

    # Process documents
    documents_path = "./service_requests"  # Update this path
    processor.process_directory(documents_path)


if __name__ == "__main__":
    main()
